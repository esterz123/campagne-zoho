# -*- coding: utf-8 -*-
"""Genere les 2 templates vendus sur mahdi-design.com/templates.html :
1. Rapport de Concurrence (DOCX, 79 EUR)
2. Kit Identite PME Industrielle (DOCX, 79 EUR)
Zero reseau, idempotent. Sortie : livrable/templates/
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "livrable", "templates")
os.makedirs(OUT, exist_ok=True)
ORANGE = RGBColor(0xFF, 0x7A, 0x1A)
GRIS = RGBColor(0x55, 0x55, 0x55)


def base_doc(titre, sous_titre):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    h = doc.add_paragraph()
    r = h.add_run(titre)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = ORANGE
    p = doc.add_paragraph()
    r = p.add_run(sous_titre)
    r.font.size = Pt(12)
    r.font.color.rgb = GRIS
    return doc


def h2(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(14)
    return p


def ph(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.color.rgb = GRIS
    r.italic = True
    return p


# ---------- 1. RAPPORT DE CONCURRENCE ----------
doc = base_doc("Rapport de Concurrence", "[Votre entreprise] vs 2 concurrents - [date]")
doc.add_paragraph()
h2(doc, "1. Terrain de jeu")
ph(doc, "[Votre secteur, votre ville, vos 3 principaux canaux d'acquisition]")
h2(doc, "2. Fiche concurrent 1 : [Nom du concurrent]")
for lbl in ["Site web / vitesse percue (lent, moyen, rapide) : ",
            "Premiere impression en 5 secondes : ",
            "Offre principale et prix affiches : ",
            "Points forts (2 max) : ",
            "Points faibles (2 max) : ",
            "Note globale /100 : "]:
    doc.add_paragraph(lbl + "[...]", style="List Bullet")
h2(doc, "3. Fiche concurrent 2 : [Nom du concurrent]")
for lbl in ["Site web / vitesse percue : ",
            "Premiere impression en 5 secondes : ",
            "Offre principale et prix affiches : ",
            "Points forts (2 max) : ",
            "Points faibles (2 max) : ",
            "Note globale /100 : "]:
    doc.add_paragraph(lbl + "[...]", style="List Bullet")
h2(doc, "4. Grille comparative")
t = doc.add_table(rows=6, cols=4)
t.style = "Light Grid Accent 1"
hdr = ["Critere", "Vous", "Concurrent 1", "Concurrent 2"]
for i, x in enumerate(hdr):
    t.rows[0].cells[i].text = x
for i, x in enumerate(["Site (modernite / vitesse)", "Clarte de l'offre", "Preuves (avis, photos, chiffres)", "Facilite de contact", "Note /100"], start=1):
    t.rows[i].cells[0].text = x
h2(doc, "5. Ou vous gagnez")
ph(doc, "[2 avantages reels de votre entreprise, prouves]")
h2(doc, "6. Ou vous perdez des clients aujourd'hui")
ph(doc, "[2 faiblesses avec le plus d'impact commercial]")
h2(doc, "7. Plan d'action priorise (30 jours)")
for i in range(1, 4):
    doc.add_paragraph("Action %d : [...] - impact estime : [...] - effort : faible/moyen" % i, style="List Number")
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Template Mahdi Design - Portfolio : mahdi-design.com")
r.font.color.rgb = GRIS
doc.save(os.path.join(OUT, "rapport_concurrence_TEMPLATE.docx"))
print("rapport_concurrence_TEMPLATE.docx OK")

# ---------- 2. KIT IDENTITE ----------
doc = base_doc("Kit Identite PME Industrielle", "Guide de style pret a remplir - secteur B2B industriel")
doc.add_paragraph()
h2(doc, "1. Positionnement en 1 phrase")
ph(doc, "[Qui vous servez] + [quel probleme vous resolvez] + [pourquoi vous plutot qu'un autre]")
h2(doc, "2. Palette de couleurs (B2B industriel)")
t = doc.add_table(rows=5, cols=4)
t.style = "Light Grid Accent 1"
for i, x in enumerate(["Role", "Couleur", "Code HEX", "Ou l'utiliser"]):
    t.rows[0].cells[i].text = x
rows = [
    ["Primaire", "Bleu acier", "#1F3A5F", "Titres, boutons, camion"],
    ["Accent", "Orange securite", "#FF7A1A", "1 element par page : CTA, chiffre cle"],
    ["Neutre sombre", "Gris anthracite", "#22262B", "Textes courants"],
    ["Neutre clair", "Gris clair", "#EDEFF2", "Fonds de section, encadres"],
]
for ri, row in enumerate(rows, start=1):
    for ci, x in enumerate(row):
        t.rows[ri].cells[ci].text = x
p = doc.add_paragraph()
r = p.add_run("Regle d'or : 80% couleurs neutres, 15% primaire, 5% accent. Jamais l'inverse.")
r.font.color.rgb = GRIS
h2(doc, "3. Typographies (gratuites, lisibles, industrielles)")
for x in ["Titres : Archivo Bold ou Barlow Condensed Bold (Google Fonts)",
          "Textes : Inter Regular, taille 16 px web / 11 pt print",
          "Interdits : Papyrus, Comic Sans, polices manuscrites, + de 2 familles"]:
    doc.add_paragraph(x, style="List Bullet")
h2(doc, "4. Logo : regles d'usage")
for x in ["Zone de protection : hauteur du 'x' du logo tout autour",
          "Version monochrome obligatoire (devis, tampon, gravure)",
          "Taille minimale : 120 px de large sur fond clair",
          "Jamais de logo sur photo chargee sans voile sombre 40%"]:
    doc.add_paragraph(x, style="List Bullet")
h2(doc, "5. Ton de voix industriel")
for x in ["Parlez resultats, pas adjectifs : 'pieces livrees en 48h' > 'service rapide'",
          "Chiffres reels partout (delais, tolerances, references)",
          "Phrases courtes. Sujet-verbe-complement. Zero jargon marketing"]:
    doc.add_paragraph(x, style="List Bullet")
h2(doc, "6. Applications a mettre a jour (checklist)")
for x in ["Signature mail (logo + tel + site)",
          "En-tete / pied de page devis et factures",
          "Page d'accueil du site (hero : 1 phrase + 1 bouton)",
          "Google Business Profile (photos + horaires reels)",
          "Vehicules et panneaux atelier"]:
    doc.add_paragraph(x, style="List Bullet")
h2(doc, "7. Exemples de structures (realisations Mahdi Design)")
for x in ["SayMyName : identite jeune, contraste fort, 1 accent par ecran",
          "OUTKA : grille stricte, typographie condensee, preuves chiffrees",
          "XANA : palette froide + accent chaud, photos terrain en fond"]:
    doc.add_paragraph(x, style="List Bullet")
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Template Mahdi Design - Portfolio : mahdi-design.com")
r.font.color.rgb = GRIS
doc.save(os.path.join(OUT, "kit_identite_pmeindustrielle_TEMPLATE.docx"))
print("kit_identite_pmeindustrielle_TEMPLATE.docx OK")
