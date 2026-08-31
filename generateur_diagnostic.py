# -*- coding: utf-8 -*-
"""GENERATEUR DE DIAGNOSTIC v1 — transforme un prospect de la file en diagnostic docx pret a livrer.
Usage : python generateur_diagnostic.py <num> [--site <domaine>]
Sortie : livrable_diagnostic/diagnostic_<num>_<entreprise>.docx
A SUPPRIMER APRES RUN : non, ce fichier est DURABLE (a commiter dans le repo)."""
import json, os, re, sys, datetime, urllib.request, urllib.parse, unicodedata

BASE = os.path.dirname(os.path.abspath(__file__))
LIV = os.path.join(BASE, 'livrable')  # 28/08 : template DANS le repo (le cloud n a pas ../livrable_diagnostic)
TEMPLATE = os.path.join(LIV, 'diagnostic_template.docx')
UA = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/126.0'}

def norm(s):
    s = unicodedata.normalize('NFD', s or '')
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn')

def fetch(url):
    try:
        req = urllib.request.Request(url, headers=UA)
        with urllib.request.urlopen(req, timeout=15) as r:
            return r.read(300000).decode('utf-8', 'ignore')
    except Exception:
        return None

def curl_site(domain):
    info = []
    html = fetch('https://' + domain) or fetch('http://' + domain)
    if not html:
        return info + ['site injoignable']
    wp = 'wp-content' in html
    gen = re.search(r'<meta name="generator" content="([^"]{0,60})"', html)
    if gen:
        info.append('CMS : %s' % gen.group(1))
    elif wp:
        info.append('CMS : WordPress (version masquee)')
    cop = re.findall(r'©\s*(?:&copy;)?\s*(\d{4})', html)
    if cop:
        info.append('copyright affiche : %s' % cop[0])
    if 'viewport' not in html:
        info.append('pas de vue mobile (viewport absent)')
    return info

def note_globale(info):
    """Note deterministe /100 a partir des signaux techniques."""
    score = 62
    for i in info:
        s = i.lower()
        if '2016' in s or '2015' in s or '2017' in s:
            score -= 12
        elif re.search(r'20(0\d|1[0-4])', s):
            score -= 10
        if 'viewport absent' in s:
            score -= 12
        if 'injoignable' in s:
            score -= 25
        if 'wordpress' in s and 'version masquee' in s:
            score -= 5
    return max(20, min(88, score))

def main():
    if len(sys.argv) < 2:
        print('Usage : python generateur_diagnostic.py <num> [--site <domaine>] [--out <chemin>]')
        sys.exit(1)
    num = int(sys.argv[1])
    site = None
    if '--site' in sys.argv:
        site = sys.argv[sys.argv.index('--site') + 1]

    camp = json.load(open(os.path.join(BASE, 'campagne_data.json'), encoding='utf-8'))
    emails = camp if isinstance(camp, list) else camp.get('emails', [])
    e = next((x for x in emails if x.get('num') == num), None)
    if not e:
        print('ERREUR : num %d introuvable dans la file' % num)
        sys.exit(1)
    entreprise = re.sub(r'^\d+ — ', '', e.get('prospect', '')).split(' (')[0]
    body = e.get('body', '')
    if not site:
        m = re.search(r'([a-z0-9-]+\.(?:fr|com|net|eu))', body)
        site = m.group(1) if m else 'le site'

    info = curl_site(site)
    note = note_globale(info)
    constat = next((ln.strip() for ln in body.split('\n') if 'logo' in ln.lower() or 'copyright' in ln.lower() or 'WordPress' in ln or 'HTTP' in ln or 'temps' in ln.lower() or 'charger' in ln.lower()), 'un site qui ne reflete plus votre niveau')
    constat = re.sub(r'^[0-9]+\.\s*', '', constat)[:220]

    # cle de remplacement
    repl = {
        '{{entreprise}}': entreprise,
        '{{date}}': datetime.date.today().strftime('%d/%m/%Y'),
        '{{constat_principal}}': constat,
        '{{note_globale}}': str(note),
        '{{element_manquant}}': 'la preuve (certifications, references, qualite de fabrication)',
        '{{action_1}}': 'Mettre a jour le copyright et le CMS (securite avant tout)',
        '{{action_2}}': 'Passer le site en HTTPS avec certificat (cadenas visible)',
        '{{action_3}}': 'Refaire l identite visuelle et le site (La Marque qui Vend)',
    }

    import docx
    d = docx.Document(TEMPLATE)

    def fix_para(p):
        if not p.text.strip():
            return
        new = p.text
        for k, v in repl.items():
            new = new.replace(k, v)
        # prix a jour (defense : meme si le template regresse)
        new = new.replace('69 €/mois', '69 €/mois').replace('à partir de 3 500 €', '3 900 € à 5 900 €').replace('3 500 €', '3 900 €')
        if new != p.text:
            for r in p.runs:
                r.text = ''
            p.runs[0].text = new if p.runs else p.add_run(new)

    for p in d.paragraphs:
        fix_para(p)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    fix_para(p)

    out = os.path.join(LIV, 'diagnostic_%d_%s.docx' % (num, re.sub(r'[^A-Za-z0-9]+', '_', entreprise)[:40]))
    if '--out' in sys.argv:  # 28/08 : chemin personnalise (livraison gratuite par le closer)
        out = sys.argv[sys.argv.index('--out') + 1]
        os.makedirs(os.path.dirname(out) or '.', exist_ok=True)
    d.save(out)
    print('OK - diagnostic genere : %s' % out)
    print('  entreprise :', entreprise, '| site :', site)
    print('  signaux :', '; '.join(info) or 'aucun')
    print('  note :', note, '/100 | constat :', constat[:100] + '...')

if __name__ == '__main__':
    main()
