# Lecture de l'etat de la campagne : reponses, revenus, docx
import json, os, re, sys

BASE = os.path.dirname(os.path.abspath(__file__))
os.chdir(BASE)

BAD = ["\u2014", "\u2013", "\u2019"]  # tiret long, tiret moyen, apostrophe typographique

def load(name):
    p = os.path.join(BASE, name)
    if not os.path.exists(p):
        return None
    try:
        return json.load(open(p, encoding="utf-8"))
    except Exception as e:
        return {"_error": str(e)}

print("=== REPONSES (campagne_state.json) ===")
cs = load("campagne_state.json")
if cs:
    sent = cs.get("sent", {})
    print("total envoyes:", len(sent))
    replied = [(k, v) for k, v in sent.items() if v.get("replied")]
    print("reponses:", len(replied))
    for k, v in replied:
        print("-", k, "| to:", v.get("to", ""), "| replied:", v.get("replied"), "| reply:", str(v.get("reply_text", v.get("reply", "")))[:300])
    # leads Gaultier / SIMI specifiquement
    for k, v in sent.items():
        to = (v.get("to") or "").lower()
        if "gaultier" in to or "simi" in to or "id-casting" in to:
            print("LEAD CIBLE:", k, json.dumps({kk: vv for kk, vv in v.items() if kk != "body"}, ensure_ascii=False)[:500])
else:
    print("campagne_state.json illisible")

print()
print("=== SUIVI REVENUS ===")
rev = load("suivi_revenus.json")
if rev is None:
    print("fichier absent (aucun paiement)")
else:
    print(json.dumps(rev, indent=1, ensure_ascii=False)[:2000])

print()
print("=== DOCX SIMI ===")
p = os.path.join(BASE, "livrable", "diagnostic_1_SIMI.docx")
print("existe:", os.path.exists(p), "| taille:", os.path.getsize(p) if os.path.exists(p) else 0)
if os.path.exists(p):
    try:
        from docx import Document
        d = Document(p)
        txt = "\n".join(par.text for par in d.paragraphs)
        bad_found = [b for b in BAD if b in txt]
        print("paragraphes:", len(d.paragraphs), "| chars:", len(txt))
        print("caracteres interdits:", bad_found if bad_found else "AUCUN (OK)")
        print("extrait:", txt[:400].replace("\n", " | "))
    except ImportError:
        # lecture zip brute
        import zipfile
        z = zipfile.ZipFile(p)
        xml = z.read("word/document.xml").decode("utf-8")
        txt = re.sub(r"<[^>]+>", " ", xml)
        bad_found = [b for b in BAD if b in txt]
        print("caracteres interdits:", bad_found if bad_found else "AUCUN (OK)")
        print("extrait:", txt[:400])

print()
print("=== REPONDEUR STATE ===")
rs = load("repondeur_state.json")
if rs:
    print(json.dumps(rs, indent=1, ensure_ascii=False)[:1500])
else:
    print("fichier absent")
