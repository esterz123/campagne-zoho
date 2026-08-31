#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GENERATEUR DE RAPPORTS DE PREUVE - rend la promesse vraie pour les 152.
=======================================================================
Chaque mail a preuve dit : "Le rapport complet est deja fait, 2 pages,
gratuit, donnez-le a votre webmaster, repondez oui." Pour que ce ne soit
pas un mensonge, ce script genere le docx AVANT la reponse, a partir des
24 faits deja mesures (constats_sites.json). Zero reseau, zero LLM.
Quand le prospect repond "oui", le docx est deja la : livraison en 15 min
au lieu de 48h.

Usage : python3 generateur_rapports.py            # tous les non-encore-envoyes
        python3 generateur_rapports.py --nums 149,165
Sortie : livrable/rapports/rapport_<num>_<domaine>.docx
"""
import os
import re
import sys
import json
import datetime

BASE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE, "livrable", "rapports")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean(t):
    return (t or "").replace("\u2019", "'").replace("\u2018", "'")


def ordre_reparations(f):
    """Sequence de reparation deduite des faits mesures (urgence client d'abord)."""
    etapes = []
    if f.get("pirate"):
        etapes.append(("1. Securite immediate", "Retirer les liens de fraude (%s) et changer les acces : un site qui diffuse des arnaques detruit votre reputation en un clic." % f["pirate"]))
    if f.get("etat") != "VIVANT":
        etapes.append(("1. Remettre le site en ligne", "Aucun client ne peut vous trouver. H\u00e9bergement ou domaine a r\u00e9activer en priorit\u00e9 absolue."))
        return etapes
    if f.get("http_seul") or not f.get("ssl_valide"):
        etapes.append(("%d. HTTPS" % (len(etapes) + 1), "Certificat en place : l'alerte rouge du navigateur disparait, Google remonte la note."))
    if f.get("parking"):
        etapes.append(("%d. Remplacer la page de parking" % (len(etapes) + 1), "Un vrai site vitrine, votre metier, vos references, un numero de telephone."))
    if f.get("mobile") is False:
        etapes.append(("%d. Version mobile" % (len(etapes) + 1), "La majorite des recherches se font sur telephone : sans affichage adapte, vous etes declasse par Google et invisible pour la moitie des visiteurs."))
    if f.get("temps_s") and f["temps_s"] > 3:
        etapes.append(("%d. Vitesse" % (len(etapes) + 1), "%.0f secondes a l'affichage aujourd'hui : compression des images et cache, objectif sous 2 s." % f["temps_s"]))
    if f.get("une_seule_page") or f.get("pages_internes", 0) <= 4:
        etapes.append(("%d. Pages metier" % (len(etapes) + 1), "Une page par savoir-faire pour que les donneurs d'ordres vous trouvent quand ils precisent leur besoin."))
    if not f.get("meta_desc"):
        etapes.append(("%d. Fiches Google" % (len(etapes) + 1), "Descriptions et titres ecrits pour vos metiers : votre resultat de recherche devient professionnel."))
    if f.get("img_total") and f.get("img_sans_alt"):
        etapes.append(("%d. Photos" % (len(etapes) + 1), "L\u00e9gendes sur les photos (atelier, machines, \u00e9quipes) : visibles dans Google Images, credibles pour un acheteur."))
    if f.get("wp_version"):
        etapes.append(("%d. Mise a jour technique" % (len(etapes) + 1), "WordPress %s : passage a jour + surveillance, c'est le Pack Serenite." % f["wp_version"]))
    if not f.get("formulaire") or not f.get("tel"):
        etapes.append(("%d. Contact" % (len(etapes) + 1), "Telephone cliquable + formulaire court : aujourd'hui un visiteur motive ne peut pas vous joindre en 10 secondes."))
    if not etapes:
        etapes.append(("1. Visibilite", "Le site est propre. La prochaine etape est la place dans Google face aux concurrents de votre departement."))
    return etapes


def rapport(f, prospect):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6); s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)

    dom = f.get("domaine", "")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RAPPORT DE PREUVE")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sl = (f.get("sonde_le") or "").strip()
    try:
        date_mesure = datetime.datetime.strptime(sl.split(" ")[0], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        date_mesure = datetime.date.today().strftime("%d/%m/%Y")
    s2.add_run("%s\nMesure en direct le %s par Mahdi Design" % (dom, date_mesure)).font.size = Pt(10)

    note = f.get("note")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = p.add_run("%s/100" % (note if note is not None else "?"))
    rn.bold = True; rn.font.size = Pt(36)
    rn.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b) if (note is not None and note < 50) else (
        RGBColor(0xe6, 0x7e, 0x22) if (note is not None and note < 75) else RGBColor(0x27, 0xae, 0x60))
    lab = doc.add_paragraph()
    lab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab.add_run("Note d'apres 24 mesures : securite, vitesse, mobile, pages, contact").font.size = Pt(9)

    doc.add_paragraph()
    h = doc.add_paragraph(); h.add_run("Ce que j'ai mesure").bold = True
    griefs = f.get("griefs") or ["Aucun defaut technique majeur releve : le travail restant est la visibilite."
                                 if (note or 0) >= 75 else "Site difficile a auditer : mesures partielles."]
    for g in griefs:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(clean(g))

    doc.add_paragraph()
    h2 = doc.add_paragraph(); h2.add_run("Dans l'ordre des reparations").bold = True
    for titre, texte in ordre_reparations(f):
        pr = doc.add_paragraph()
        pr.add_run(titre + " : ").bold = True
        pr.add_run(texte)

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.add_run("La suite, sans engagement : ").bold = True
    fin.add_run("ce rapport est a vous, donnez-le a votre webmaster. Si vous voulez que je m'en occupe : "
                "Pack Serenite 69 EUR/mois (securite, mises a jour, hebergement, sans engagement) ou refonte "
                "complete (devis 3900 EUR, a la rentree 2900 EUR).")
    sig = doc.add_paragraph()
    sig.add_run("\nMahdi - Portfolio : mahdi-design.com - contact@mahdi-design.com").font.size = Pt(9)
    return doc


def main():
    nums_only = None
    for i, a in enumerate(sys.argv):
        if a == "--nums" and i + 1 < len(sys.argv):
            nums_only = set(x.strip() for x in sys.argv[i + 1].split(","))
    os.makedirs(OUT_DIR, exist_ok=True)
    preuves = json.load(open(os.path.join(BASE, "constats_sites.json"), encoding="utf-8"))
    data = json.load(open(os.path.join(BASE, "campagne_data.json"), encoding="utf-8"))
    state = json.load(open(os.path.join(BASE, "campagne_state.json"), encoding="utf-8"))
    sent = set(str(k) for k in state.get("sent", {}))
    par_num = {str(r.get("num")): r for r in data}

    faits = 0
    for num, f in preuves.items():
        if num in sent:
            continue  # les deja-envoyes auront leur rapport a la reponse (GO2)
        if nums_only and num not in nums_only:
            continue
        r = par_num.get(num, {})
        fn = os.path.join(OUT_DIR, "rapport_%s_%s.docx" % (num, re.sub(r"[^a-z0-9.-]", "_", f.get("domaine", "x"))))
        if os.path.exists(fn):
            continue
        rapport(f, r.get("prospect", "")).save(fn)
        faits += 1
    print("rapports generes: %d -> %s" % (faits, OUT_DIR))
    total = len([x for x in os.listdir(OUT_DIR) if x.endswith(".docx")])
    print("total dispo: %d" % total)


if __name__ == "__main__":
    sys.exit(main())
